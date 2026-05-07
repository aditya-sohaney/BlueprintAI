"""Automated PDF/DOCX/Excel report generation."""

import json
from pathlib import Path
from datetime import datetime

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt

from core.database import DrawingDatabase


def generate_excel_report(df: pd.DataFrame = None, output_path: str = None) -> str:
    """Generate a comprehensive Excel report with multiple sheets.

    Args:
        df: DataFrame to use (if None, loads from database).
        output_path: Output file path.

    Returns:
        Path to the generated Excel file.
    """
    if df is None:
        db = DrawingDatabase()
        df = db.export_to_dataframe()
        db.close()

    if output_path is None:
        output_path = (Path(__file__).parent.parent / "data" / "exports" /
                       f"adot_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    with pd.ExcelWriter(str(output_path), engine="openpyxl") as writer:
        # Sheet 1: All Drawings
        df.to_excel(writer, sheet_name="All Drawings", index=False)

        # Sheet 2: Summary by Firm
        if "firm" in df.columns:
            firm_summary = df.groupby("firm").agg({
                "page_number": "count",
                "design_duration_days": "mean",
                "rfc_duration_days": "mean",
            }).round(1)
            firm_summary.columns = ["Drawing Count", "Avg Design Days", "Avg RFC Days"]
            firm_summary.to_excel(writer, sheet_name="By Firm")

        # Sheet 3: Summary by Engineer
        if "engineer_stamp_name" in df.columns:
            eng_summary = df.groupby("engineer_stamp_name").agg({
                "page_number": "count",
                "design_duration_days": "mean",
                "rfc_duration_days": "mean",
            }).round(1)
            eng_summary.columns = ["Drawing Count", "Avg Design Days", "Avg RFC Days"]
            eng_summary.to_excel(writer, sheet_name="By Engineer")

        # Sheet 4: Date Analysis
        date_cols = ["initial_date", "final_date", "rfc_date"]
        for col in date_cols:
            if col in df.columns:
                df[col] = pd.to_datetime(df[col], format="%m/%d/%Y", errors="coerce")

        if all(c in df.columns for c in date_cols):
            date_df = df[["rw_number", "drawing_title"] + date_cols +
                         ["design_duration_days", "rfc_duration_days"]].copy()
            date_df.to_excel(writer, sheet_name="Date Analysis", index=False)

    print(f"Excel report saved to {output_path}")
    return str(output_path)


def generate_docx_report(df: pd.DataFrame = None, output_path: str = None) -> str:
    """Generate a Word document report.

    Args:
        df: DataFrame to use (if None, loads from database).
        output_path: Output file path.

    Returns:
        Path to the generated DOCX file.
    """
    if df is None:
        db = DrawingDatabase()
        df = db.export_to_dataframe()
        db.close()

    if output_path is None:
        output_path = (Path(__file__).parent.parent / "data" / "exports" /
                       f"adot_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Title
    doc.add_heading("ADOT Engineering Drawing Extraction Report", 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Total Drawings Analyzed: {len(df)}")

    # Summary Section
    doc.add_heading("Summary Statistics", level=1)

    if "firm" in df.columns:
        doc.add_heading("Firm Distribution", level=2)
        for firm, count in df["firm"].value_counts().items():
            doc.add_paragraph(f"{firm}: {count} drawings", style="List Bullet")

    if "engineer_stamp_name" in df.columns:
        doc.add_heading("Engineer Distribution", level=2)
        for eng, count in df["engineer_stamp_name"].dropna().value_counts().items():
            doc.add_paragraph(f"{eng}: {count} drawings", style="List Bullet")

    if "design_duration_days" in df.columns:
        dur = df["design_duration_days"].dropna()
        if len(dur) > 0:
            doc.add_heading("Design Duration", level=2)
            doc.add_paragraph(f"Average: {dur.mean():.1f} days")
            doc.add_paragraph(f"Median: {dur.median():.1f} days")
            doc.add_paragraph(f"Range: {dur.min():.0f} - {dur.max():.0f} days")

    # Per-Drawing Table
    doc.add_heading("Drawing Details", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    headers = ["RW Number", "Drawing Title", "Firm", "Engineer", "RFC Date"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    for _, row in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row.get("rw_number", ""))
        cells[1].text = str(row.get("drawing_title", ""))[:60]
        cells[2].text = str(row.get("firm", ""))
        cells[3].text = str(row.get("engineer_stamp_name", ""))
        cells[4].text = str(row.get("rfc_date", ""))

    doc.save(str(output_path))
    print(f"DOCX report saved to {output_path}")
    return str(output_path)
